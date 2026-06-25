import os
import sys
import subprocess
import datetime
from pathlib import Path
import json

# Try importing docx, handle gracefully if not installed yet in sandbox
try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

TEMPLATE_DIR = Path("templates")
CONFIG_PATH = Path("config.json")

from database import load_config, save_config

def bootstrap_templates():
    """Create default templates if they do not exist."""
    if not docx:
        print("python-docx is not installed. Skipping template bootstrapping.")
        return

    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    config = load_config()
    
    instansi = config.get("instansi", {})
    desa = instansi.get("desa", "Kampokku Jaya").upper()
    kec = instansi.get("kecamatan", "Pakkampong").upper()
    kab = instansi.get("kabupaten", "Limpo").upper()
    prov = instansi.get("provinsi", "Limpo Toddang").upper()

    templates_to_create = {
        "Domisili.docx": {
            "title": "SURAT KETERANGAN DOMISILI",
            "content": "Yang bertanda tangan di bawah ini Kepala Desa {{DESA}}, Kecamatan {{KECAMATAN}}, Kabupaten {{KABUPATEN}}, Provinsi {{PROVINSI}}, dengan ini menerangkan bahwa penduduk di bawah ini:"
        },
        "SKTM.docx": {
            "title": "SURAT KETERANGAN TIDAK MAMPU",
            "content": "Yang bertanda tangan di bawah ini Kepala Desa {{DESA}}, Kecamatan {{KECAMATAN}}, Kabupaten {{KABUPATEN}}, Provinsi {{PROVINSI}}, menerangkan dengan sebenarnya bahwa:"
        },
        "Usaha.docx": {
            "title": "SURAT KETERANGAN USAHA",
            "content": "Yang bertanda tangan di bawah ini Kepala Desa {{DESA}}, Kecamatan {{KECAMATAN}}, Kabupaten {{KABUPATEN}}, Provinsi {{PROVINSI}}, dengan ini menerangkan bahwa:"
        },
        "BelumMenikah.docx": {
            "title": "SURAT KETERANGAN BELUM MENIKAH",
            "content": "Yang bertanda tangan di bawah ini Kepala Desa {{DESA}}, Kecamatan {{KECAMATAN}}, Kabupaten {{KABUPATEN}}, Provinsi {{PROVINSI}}, menerangkan dengan sebenarnya bahwa:"
        }
    }

    for filename, t_info in templates_to_create.items():
        file_path = TEMPLATE_DIR / filename
        if file_path.exists():
            continue

        doc = docx.Document()
        
        # Set standard margins (1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header Instansi
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_prov = header_p.add_run(f"PEMERINTAH PROVINSI {prov}\n")
        run_prov.font.name = "Arial"
        run_prov.font.size = Pt(12)
        run_prov.font.bold = True
        
        run_kab = header_p.add_run(f"PEMERINTAH KABUPATEN {kab}\n")
        run_kab.font.name = "Arial"
        run_kab.font.size = Pt(14)
        run_kab.font.bold = True
        
        run_kec = header_p.add_run(f"KECAMATAN {kec}\n")
        run_kec.font.name = "Arial"
        run_kec.font.size = Pt(14)
        run_kec.font.bold = True
        
        run_desa = header_p.add_run(f"KANTOR KEPALA DESA {desa}\n")
        run_desa.font.name = "Arial"
        run_desa.font.size = Pt(16)
        run_desa.font.bold = True

        run_detail = header_p.add_run("Alamat: {{ALAMAT_INSTANSI}}  Telp: {{TELEPON_INSTANSI}}\n")
        run_detail.font.name = "Arial"
        run_detail.font.size = Pt(10)
        run_detail.italic = True

        # Horizontal line (border)
        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_line = p_line.add_run("__________________________________________________________________")
        run_line.font.bold = True
        
        doc.add_paragraph() # Spacer

        # Letter Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = title_p.add_run(f"{t_info['title']}\n")
        run_title.font.name = "Arial"
        run_title.font.size = Pt(14)
        run_title.font.bold = True
        run_title.font.underline = True

        run_no = title_p.add_run("Nomor: {{NOMOR_SURAT}}")
        run_no.font.name = "Arial"
        run_no.font.size = Pt(11)

        doc.add_paragraph() # Spacer

        # Content Opening
        p_open = doc.add_paragraph()
        p_open.paragraph_format.line_spacing = 1.15
        run_open = p_open.add_run(t_info["content"])
        run_open.font.name = "Arial"
        run_open.font.size = Pt(11)

        doc.add_paragraph() # Spacer

        # Resident Data Grid (Table format for neat layout)
        table = doc.add_table(rows=0, cols=3)
        table.autofit = False
        
        # Set widths for nice alignment
        col_widths = [Inches(2.0), Inches(0.2), Inches(4.3)]
        
        fields = [
            ("Nama Lengkap", "{{NAMA}}"),
            ("NIK", "{{NIK}}"),
            ("No. Kartu Keluarga", "{{KK}}"),
            ("Tempat / Tanggal Lahir", "{{TEMPAT_LAHIR}}, {{TANGGAL_LAHIR}}"),
            ("Jenis Kelamin", "{{JENIS_KELAMIN}}"),
            ("Agama", "{{AGAMA}}"),
            ("Pekerjaan", "{{PEKERJAAN}}"),
            ("Status Perkawinan", "{{STATUS_PERKAWINAN}}"),
            ("Kewarganegaraan", "{{KEWARGANEGARAAN}}"),
            ("Alamat Lengkap", "{{ALAMAT}} RT {{RT}} / RW {{RW}}")
        ]

        for label, val in fields:
            row = table.add_row()
            
            # Label
            cell_label = row.cells[0]
            cell_label.width = col_widths[0]
            p = cell_label.paragraphs[0]
            run = p.add_run(label)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            
            # Colon
            cell_colon = row.cells[1]
            cell_colon.width = col_widths[1]
            p = cell_colon.paragraphs[0]
            run = p.add_run(":")
            run.font.name = "Arial"
            run.font.size = Pt(11)

            # Value
            cell_val = row.cells[2]
            cell_val.width = col_widths[2]
            p = cell_val.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(11)

        doc.add_paragraph() # Spacer

        # Extra description depending on letter type
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.line_spacing = 1.15
        
        if filename == "Domisili.docx":
            run_desc = p_desc.add_run(
                "Orang tersebut di atas adalah benar penduduk domisili Desa {{DESA}}, "
                "Kecamatan {{KECAMATAN}}, Kabupaten {{KABUPATEN}} dan bertempat tinggal pada alamat tersebut."
            )
        elif filename == "SKTM.docx":
            run_desc = p_desc.add_run(
                "Berdasarkan keterangan yang ada pada kami dan sepengetahuan kami, yang bersangkutan benar-benar "
                "tergolong keluarga yang tidak mampu ekonomi dan layak untuk mendapatkan bantuan/fasilitas keringanan."
            )
        elif filename == "Usaha.docx":
            run_desc = p_desc.add_run(
                "Berdasarkan keterangan yang bersangkutan, nama tersebut di atas benar-benar memiliki kegiatan usaha "
                "di bidang {{PEKERJAAN}} yang berlokasi di Desa {{DESA}}."
            )
        else: # BelumMenikah.docx
            run_desc = p_desc.add_run(
                "Berdasarkan catatan kependudukan kami, nama tersebut di atas benar-benar belum pernah menikah "
                "dengan siapapun dan berstatus lajang/belum kawin."
            )
            
        run_desc.font.name = "Arial"
        run_desc.font.size = Pt(11)

        doc.add_paragraph()

        # Closing paragraph
        p_close = doc.add_paragraph()
        run_close = p_close.add_run(
            "Demikian surat keterangan ini dibuat dengan sebenarnya untuk dapat dipergunakan sebagaimana mestinya."
        )
        run_close.font.name = "Arial"
        run_close.font.size = Pt(11)

        doc.add_paragraph() # Spacer
        doc.add_paragraph()

        # Signature Table (Aligns signature to the right)
        sig_table = doc.add_table(rows=4, cols=2)
        
        # Hide borders (default in python-docx table is no border, but let's be sure)
        sig_widths = [Inches(3.5), Inches(3.0)]
        
        # Row 0: Date
        cell_date = sig_table.rows[0].cells[1]
        cell_date.width = sig_widths[1]
        p = cell_date.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("{{DESA}}, {{TANGGAL_SEKARANG}}")
        run.font.name = "Arial"
        run.font.size = Pt(11)

        # Row 1: Title/Jabatan
        cell_jab = sig_table.rows[1].cells[1]
        cell_jab.width = sig_widths[1]
        p = cell_jab.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("{{JABATAN_KADES}} {{DESA}}")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.bold = True

        # Row 2: Spacer (Spasir Tanda Tangan)
        cell_space = sig_table.rows[2].cells[1]
        cell_space.width = sig_widths[1]
        p = cell_space.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n\n\n") # Signature space

        # Row 3: Name and NIP
        cell_name = sig_table.rows[3].cells[1]
        cell_name.width = sig_widths[1]
        p = cell_name.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_name = p.add_run("{{NAMA_KADES}}\n")
        run_name.font.name = "Arial"
        run_name.font.size = Pt(11)
        run_name.font.bold = True
        run_name.font.underline = True
        
        run_nip = p.add_run("NIP. {{NIP_KADES}}")
        run_nip.font.name = "Arial"
        run_nip.font.size = Pt(10)

        # Save document
        doc.save(str(file_path))
        print(f"Created template: {filename}")


def replace_text_in_paragraph(paragraph, placeholders):
    """Replace placeholders in docx paragraph while preserving styling runs where possible."""
    text = paragraph.text
    changed = False
    
    for key, value in placeholders.items():
        tag = f"{{{{{key}}}}}"
        if tag in text:
            # First check if the tag fits entirely inside a single run
            for run in paragraph.runs:
                if tag in run.text:
                    run.text = run.text.replace(tag, str(value))
                    changed = True
            
            # If the tag is split across runs, we fall back to paragraph-level replacement
            # which might reset local styling within this paragraph
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, str(value))
                changed = True
                
    return changed


def generate_document(template_filename, resident_data, letter_number):
    """
    Load a template, replace all placeholders, and save the output DOCX.
    Returns: output_docx_path, error_message
    """
    if not docx:
        return None, "python-docx tidak terpasang."

    template_path = TEMPLATE_DIR / template_filename
    if not template_path.exists():
        return None, f"Berkas template '{template_filename}' tidak ditemukan."

    config = load_config()
    instansi = config.get("instansi", {})
    signatory = config.get("penandatangan", {})
    storage = config.get("penyimpanan", {})

    docx_output_dir = Path(storage.get("folder_docx", "output/docx"))
    docx_output_dir.mkdir(parents=True, exist_ok=True)

    # 1. Prepare Placeholders Mapping
    today = datetime.date.today().strftime("%d-%m-%Y")
    
    placeholders = {
        # Instansi Profile
        "DESA": instansi.get("desa", "Kampokku Jaya"),
        "KECAMATAN": instansi.get("kecamatan", "Pakkampong"),
        "KABUPATEN": instansi.get("kabupaten", "Limpo"),
        "PROVINSI": instansi.get("provinsi", "Limpo Toddang"),
        "ALAMAT_INSTANSI": instansi.get("alamat", ""),
        "TELEPON_INSTANSI": instansi.get("telepon", ""),

        # Letter Meta
        "NOMOR_SURAT": letter_number,
        "TANGGAL_SEKARANG": today,

        # Signatory
        "NAMA_KADES": signatory.get("nama", ""),
        "NIP_KADES": signatory.get("nip", ""),
        "JABATAN_KADES": signatory.get("jabatan", "Kepala Desa"),

        # Resident Data
        "NIK": resident_data.get("nik", ""),
        "KK": resident_data.get("kk", ""),
        "NAMA": resident_data.get("nama", ""),
        "TEMPAT_LAHIR": resident_data.get("tempat_lahir", ""),
        "TANGGAL_LAHIR": resident_data.get("tanggal_lahir", ""),
        "JENIS_KELAMIN": resident_data.get("jenis_kelamin", ""),
        "AGAMA": resident_data.get("agama", ""),
        "STATUS_PERKAWINAN": resident_data.get("status_perkawinan", ""),
        "PEKERJAAN": resident_data.get("pekerjaan", ""),
        "ALAMAT": resident_data.get("alamat", ""),
        "RT": resident_data.get("rt", ""),
        "RW": resident_data.get("rw", ""),
        "KEWARGANEGARAAN": resident_data.get("kewarganegaraan", "WNI")
    }

    try:
        # Load template
        doc = docx.Document(str(template_path))

        # Replace in all paragraphs
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, placeholders)

        # Replace in all tables (resident grid & signature)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, placeholders)

        # Clean filename: replace / with _ to prevent directory errors
        clean_number = letter_number.replace("/", "_")
        output_filename = f"{template_path.stem}_{resident_data.get('nik')}_{clean_number}.docx"
        output_path = docx_output_dir / output_filename
        
        doc.save(str(output_path))
        return output_path, None

    except Exception as e:
        return None, f"Gagal membuat dokumen Word: {str(e)}"


def convert_to_pdf(docx_path, libreoffice_path=None):
    """
    Convert DOCX to PDF using LibreOffice (cross-platform) or Microsoft Word COM (Windows only).
    libreoffice_path: Custom path to soffice executable (optional)
    Returns: output_pdf_path, error_message
    """
    config = load_config()
    storage = config.get("penyimpanan", {})
    pdf_output_dir = Path(storage.get("folder_pdf", "output/pdf"))
    pdf_output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = Path(docx_path)
    output_pdf_path = pdf_output_dir / f"{docx_path.stem}.pdf"

    # --- Windows: Try Microsoft Word COM first ---
    if sys.platform == "win32" and not libreoffice_path:
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # COM requires absolute paths
            abs_docx = docx_path.resolve()
            abs_pdf = output_pdf_path.resolve()
            
            doc = word.Documents.Open(str(abs_docx))
            # 17 is the constant for saving as PDF
            doc.SaveAs(str(abs_pdf), FileFormat=17)
            doc.Close()
            word.Quit()
            
            if abs_pdf.exists():
                return abs_pdf, None
        except Exception as e:
            print("Windows MS Word PDF conversion failed, falling back to LibreOffice:", str(e))

    # --- Fallback: LibreOffice Headless ---
    # Search paths for LibreOffice soffice
    soffice_exec = "soffice"
    if libreoffice_path and Path(libreoffice_path).exists():
        soffice_exec = libreoffice_path
    else:
        # Standard system locations
        if sys.platform == "win32":
            possible_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"LibreOfficePortable\App\libreoffice\program\soffice.exe" # bundled path helper
            ]
            for p in possible_paths:
                if Path(p).exists():
                    soffice_exec = p
                    break
        elif sys.platform == "darwin":
            soffice_exec = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        else:
            # Linux: check standard 'soffice' or 'libreoffice' on PATH
            soffice_exec = "libreoffice"

    try:
        # LibreOffice command line convert
        # libreoffice --headless --convert-to pdf --outdir <outdir> <docx_path>
        cmd = [
            soffice_exec,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(pdf_output_dir),
            str(docx_path)
        ]
        
        # Run process
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=30)
        
        if result.returncode == 0 and output_pdf_path.exists():
            return output_pdf_path, None
        else:
            # Check if libreoffice command is missing
            err_msg = result.stderr if result.stderr else "Gagal menjalankan LibreOffice."
            if "not found" in err_msg or "No such file" in err_msg:
                return None, "LibreOffice tidak terpasang di sistem. Silakan pasang LibreOffice atau tentukan jalurnya di Pengaturan."
            return None, f"Konversi PDF gagal: {err_msg}"
            
    except subprocess.TimeoutExpired:
        return None, "Konversi PDF waktu habis (timeout)."
    except Exception as e:
        return None, f"Kesalahan konversi PDF: {str(e)}"


def print_document(file_path, printer_name=None):
    """
    Print a PDF or DOCX file using system default or selected printer.
    Returns: success (bool), error_message
    """
    file_path = Path(file_path)
    if not file_path.exists():
        return False, "Berkas tidak ditemukan."

    try:
        if sys.platform == "win32":
            import win32api
            import win32print
            
            # Print DOCX or PDF using shell execution
            # "print" verb uses default printer, "printto" allows specific printer
            if printer_name:
                win32api.ShellExecute(0, "printto", str(file_path), f'"{printer_name}"', ".", 0)
            else:
                win32api.ShellExecute(0, "print", str(file_path), None, ".", 0)
            return True, None
        
        elif sys.platform == "linux":
            # Call lp or lpr to print
            cmd = ["lp"]
            if printer_name:
                cmd.extend(["-d", printer_name])
            cmd.append(str(file_path))
            
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if result.returncode == 0:
                return True, None
            else:
                return False, f"Gagal mencetak: {result.stderr}"
        
        else:
            # macOS lpr command
            cmd = ["lpr"]
            if printer_name:
                cmd.extend(["-P", printer_name])
            cmd.append(str(file_path))
            
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if result.returncode == 0:
                return True, None
            else:
                return False, f"Gagal mencetak: {result.stderr}"
                
    except Exception as e:
        return False, str(e)


def get_available_printers():
    """Get list of installed printer names on the OS."""
    printers = []
    try:
        if sys.platform == "win32":
            import win32print
            flags = win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS
            for printer in win32print.EnumPrinters(flags, None, 1):
                printers.append(printer[2])
        elif sys.platform == "linux" or sys.platform == "darwin":
            # parse lpstat output
            result = subprocess.run(["lpstat", "-a"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if result.returncode == 0:
                for line in result.stdout.splitlines():
                    if line.strip():
                        printers.append(line.split()[0])
    except Exception as e:
        print("Error fetching printers:", str(e))
    return printers
